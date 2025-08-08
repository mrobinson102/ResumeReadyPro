# streamlit_app.py
# ResumeReadyPro - Hybrid (Offline + GPT-ready)
# ------------------------------------------------
# Features:
# - Local auth (login/register/change password) with JSON storage
# - Onboarding wizard for first-time users
# - Generate Summary (Offline mock or GPT when enabled)
# - Upload Resume -> extract text + generate interview questions (Offline/GPT)
# - Prompt Lab (templates + tone controls)
# - Job Fit & Salary Alignment Analyzer (weighted: required vs nice-to-have, seniority, years)
# - Admin Dashboard with metrics + JSON/CSV export
# - Export to TXT/PDF/DOCX
#
# Toggle GPT on by setting USE_GPT=True and providing OPENAI_API_KEY in your env.

import os, json, re, io, hashlib
from datetime import datetime
from typing import Dict, List, Tuple
import streamlit as st

# Optional deps (gracefully degrade)
try:
    import pandas as pd
except Exception:
    pd = None

try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document as DocxDocument  # python-docx
except Exception:
    DocxDocument = None

try:
    from fpdf import FPDF
except Exception:
    FPDF = None

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
except Exception:
    plt = None

# ---------------- GPT toggle ----------------
USE_GPT = False
OPENAI_MODEL = "gpt-4o-mini"
try:
    import openai
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
    if USE_GPT and OPENAI_API_KEY:
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)  # type: ignore
    else:
        client = None
except Exception:
    client = None

APP_TITLE = "ResumeReadyPro (Hybrid: Offline + GPT-ready)"
USERS_DB = "user_data.json"

# ---------------- Storage helpers ----------------
def _ensure_db():
    if not os.path.exists(USERS_DB):
        with open(USERS_DB, "w") as f:
            json.dump({"users": {}, "metrics": {"summaries":0,"resumes":0,"questions":0,"gap_analyses":0}}, f, indent=2)

def load_db() -> Dict:
    _ensure_db()
    try:
        with open(USERS_DB, "r") as f:
            return json.load(f)
    except Exception:
        return {"users": {}, "metrics": {"summaries":0,"resumes":0,"questions":0,"gap_analyses":0}}

def save_db(db: Dict):
    with open(USERS_DB, "w") as f:
        json.dump(db, f, indent=2)

DB = load_db()

# ---------------- Auth ----------------
def hash_pw(pw:str) -> str:
    return hashlib.sha256(pw.encode("utf-8")).hexdigest()

def auth_seed_admin():
    users = DB.get("users", {})
    if "admin" not in users:
        users["admin"] = {
            "name": "Admin User",
            "pw": hash_pw("adminpass"),
            "created": datetime.utcnow().isoformat(),
            "reset_token": ""
        }
        DB["users"] = users
        save_db(DB)

def authenticate(username, password) -> bool:
    u = DB.get("users", {}).get(username)
    if not u:
        return False
    return hash_pw(password) == u.get("pw")

def register_user(username, name, password) -> Tuple[bool,str]:
    username = (username or "").strip().lower()
    if not username or not password:
        return False, "Username and password are required."
    if username in DB.get("users", {}):
        return False, "User already exists."
    DB["users"][username] = {
        "name": name or username, "pw": hash_pw(password),
        "created": datetime.utcnow().isoformat(),
        "reset_token": ""
    }
    save_db(DB)
    return True, "User registered."

def change_password(username, old, new) -> Tuple[bool,str]:
    if not authenticate(username, old):
        return False, "Old password incorrect."
    DB["users"][username]["pw"] = hash_pw(new)
    save_db(DB)
    return True, "Password changed."

def create_reset_token(username) -> Tuple[bool,str]:
    users = DB.get("users", {})
    if username not in users:
        return False, "No such user."
    token = hashlib.sha256(f"{username}{datetime.utcnow().isoformat()}".encode()).hexdigest()[:12]
    users[username]["reset_token"] = token
    DB["users"] = users
    save_db(DB)
    return True, token

def reset_password_with_token(username, token, new_pw) -> Tuple[bool,str]:
    users = DB.get("users", {})
    u = users.get(username)
    if not u:
        return False, "No such user."
    if token and u.get("reset_token") == token:
        u["pw"] = hash_pw(new_pw)
        u["reset_token"] = ""
        save_db(DB)
        return True, "Password reset successful."
    return False, "Invalid token."

auth_seed_admin()

# ---------------- Session ----------------
def ensure_session():
    if "auth" not in st.session_state:
        st.session_state.auth = {"logged_in": False, "user": None}
    if "onboarded" not in st.session_state:
        st.session_state.onboarded = False
    if "metrics" not in st.session_state:
        st.session_state.metrics = DB.get("metrics", {"summaries":0,"resumes":0,"questions":0,"gap_analyses":0})
ensure_session()

# ---------------- Onboarding ----------------
def onboarding_wizard():
    with st.expander("👋 Start here — Quick Onboarding", expanded=not st.session_state.onboarded):
        st.write("""
**Welcome to ResumeReadyPro!**  
Here’s how to get value fast:

1) **Generate Summary** — create a clean, ATS-friendly intro  
2) **Upload Resume** — extract text and practice interview questions  
3) **Job Fit & Salary** — paste a JD, get a fit score, gaps, and salary alignment  
4) **Prompt Lab** — templates & tone control  
5) **Admin Dashboard** — see usage and export data

> This build runs **offline** by default. Flip to **GPT mode** later with `USE_GPT=True` + API key.
""")
        if st.button("Got it — hide this"):
            st.session_state.onboarded = True

# ---------------- Common text utilities ----------------
TECH_KEYWORDS = {
    "python","r","sql","excel","tableau","power bi","pandas","numpy","sklearn","scikit-learn",
    "tensorflow","pytorch","spark","hadoop","airflow","dbt","git","github","docker","kubernetes",
    "aws","azure","gcp","bigquery","snowflake","databricks","redshift","postgres","mysql","graphql",
    "rest","fastapi","flask","django","react","typescript","javascript","bash","linux","terraform",
    "ansible","mlops","nlp","llm","security","nist","rmf","fedramp","stigs","clearance"
}
SOFT_SKILLS = {
    "leadership","communication","collaboration","mentoring","stakeholder","ownership",
    "problem solving","critical thinking","presentation","planning","prioritization"
}

def normalize(text:str) -> str:
    return re.sub(r"\s+"," ", (text or "").lower()).strip()

def extract_keywords(text:str) -> List[str]:
    t = normalize(text)
    found = set()
    for kw in TECH_KEYWORDS.union(SOFT_SKILLS):
        if kw in t:
            found.add(kw)
    return sorted(found)

# ---------------- Salary helpers ----------------
SALARY_BANDS = {
    "data analyst": (65000, 85000, 110000),
    "data scientist": (100000, 135000, 175000),
    "ml engineer": (120000, 160000, 210000),
    "software engineer": (100000, 140000, 190000),
    "devops engineer": (110000, 145000, 185000),
    "cloud engineer": (115000, 150000, 200000),
    "product manager": (110000, 145000, 190000),
    "it project manager": (95000, 120000, 150000),
    "security engineer": (115000, 155000, 210000),
    "solutions architect": (125000, 165000, 220000),
}
LOCATION_MULTIPLIER = {"remote": 1.0, "low-cost": 0.9, "standard": 1.0, "high-cost": 1.15}

def estimate_salary_band(role: str, location_level: str = "standard"):
    role_key = (role or "").strip().lower()
    base = SALARY_BANDS.get(role_key)
    mult = LOCATION_MULTIPLIER.get(location_level or "standard", 1.0)
    if not base: return None
    lo, mid, hi = base
    return (int(lo*mult), int(mid*mult), int(hi*mult))

def compare_salary(expected: int, band: tuple[int,int,int]) -> dict:
    lo, mid, hi = band
    if expected < lo:
        status = "Below Market"
        note = f"Your expectation (${expected:,}) is **below** market (${lo:,}–${hi:,}). Consider asking closer to mid."
    elif expected > hi:
        status = "Above Market"
        note = f"Your expectation (${expected:,}) is **above** market (${lo:,}–${hi:,}). Consider moderating by 10–20% or justify scope/impact."
    else:
        status = "Within Market"
        note = f"Your expectation (${expected:,}) is **within** market (${lo:,}–${hi:,})."
    return {"status": status, "note": note, "band_low": lo, "band_mid": mid, "band_high": hi}

# ---------------- Extractors ----------------
def extract_text_pdf(uploaded_file) -> str:
    if not PdfReader: return ""
    try:
        reader = PdfReader(uploaded_file)
        return "\n".join([(p.extract_text() or "") for p in reader.pages])
    except Exception:
        return ""

def extract_text_docx(uploaded_file) -> str:
    if not DocxDocument: return ""
    try:
        doc = DocxDocument(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception:
        return ""

def extract_text_generic(uploaded_file) -> str:
    name = (uploaded_file.name or "").lower()
    if name.endswith(".pdf"):
        return extract_text_pdf(uploaded_file)
    if name.endswith(".docx"):
        return extract_text_docx(uploaded_file)
    try:
        return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""

# ---------------- Exports ----------------
def export_pdf(text: str) -> bytes:
    if not FPDF:
        return text.encode("utf-8")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        pdf.multi_cell(0, 8, line)
    return pdf.output(dest="S").encode("latin-1")

def export_docx(text: str) -> bytes:
    if not DocxDocument:
        return text.encode("utf-8")
    buf = io.BytesIO()
    doc = DocxDocument()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ---------------- Offline generators ----------------
def summary_offline(full_name:str, role:str, experience:str, skills:str) -> str:
    exp_kw = extract_keywords(experience)
    skill_kw = extract_keywords(skills)
    highlights = list(dict.fromkeys(exp_kw + skill_kw))[:10]
    bullets = "\n".join([f"- Experience with **{kw}**" for kw in highlights]) if highlights else "- Strong fundamentals and rapid learning"
    return f"""**{full_name or 'Candidate'}** — {role or 'Target Role'}

Collaborative professional delivering reliable solutions in fast-paced environments, translating requirements into measurable outcomes.

**Highlights**
{bullets}

**Value**
- Communicates clearly with technical & non-technical stakeholders
- Ownership mindset: deliver, measure, iterate
- Continuous improvement and mentoring culture
"""

def questions_offline(text:str, qtype:str, count:int) -> List[str]:
    base = extract_keywords(text) or ["teamwork","problem solving","ownership","python"]
    out = []
    for i in range(count):
        kw = base[i % len(base)]
        if qtype == "Behavioral":
            out.append(f"Tell me about a time you demonstrated {kw}. What was the context, actions, and outcome?")
        elif qtype == "Technical":
            out.append(f"How would you use {kw} to design or troubleshoot a real-world system? Be specific.")
        else:
            if i % 2 == 0:
                out.append(f"Walk me through a project where {kw} was critical. What design decisions did you make?")
            else:
                out.append(f"Describe a challenge involving {kw}. How did you collaborate and ensure delivery?")
    return out

# ---------------- GPT helpers (only used when USE_GPT=True) ----------------
def gpt_chat(prompt: str) -> str:
    if not (USE_GPT and client):
        return f"(Offline mock)\n\n{prompt[:300]}\n\n— This would be replaced by GPT output when you enable billing."
    try:
        resp = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[{"role":"user","content":prompt}]
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return f"(GPT error) {e}"

# ---------------- Prompt templates (offline + GPT-ready) ----------------
PROMPT_TEMPLATES = {
    "Concise Tech Summary": """Write a 3–4 sentence, ATS-friendly professional summary.
Target role: {role}
Key strengths: {skills}
Experience context: {experience}
Tone: {tone}. Keep it concise and metrics-aware.""",

    "Impact Bullets (5)": """Create 5 resume bullets focused on measurable outcomes.
Role: {role}
Inputs: {experience}
Skills: {skills}
Tone: {tone}. Use strong verbs and numbers (e.g., reduced cost by X%).""",

    "Internship/Co-op Section": """Generate a section titled 'Internship & Co-op Experience' with 3–6 bullets.
Context: {experience}
Skills: {skills}
Tone: {tone}. Include company names (generic if missing), dates (approximate ok), and technologies.""",

    "Projects – Categorized": """Categorize projects into Internship, Academic, and Personal.
Projects text: {experience}
Skills: {skills}
Tone: {tone}. Provide 2–4 bullets per category (if relevant) with tech + result.""",
}
PROMPT_TONES = ["Neutral", "Confident", "Executive", "Friendly", "Direct", "Data-driven"]

# ---------------- JD parsing & weighting helpers ----------------
RE_REQUIRED = re.compile(r"(must[-\s]?have|required|requirements|we require|you must)", re.I)
RE_NICE     = re.compile(r"(nice[-\s]?to[-\s]?have|preferred|bonus|plus)", re.I)
RE_YEARS    = re.compile(r"(\d+)\+?\s+(years|yrs)", re.I)
SENIOR_SIGNALS = {
    "senior","lead","principal","staff","manager","management","architect","strategy","roadmap","mentoring","mentorship"
}
JUNIOR_SIGNALS = {"junior","assoc","associate","entry","new grad","intern"}

def infer_seniority(text: str) -> str:
    t = normalize(text)
    has_senior = any(s in t for s in SENIOR_SIGNALS)
    has_junior = any(j in t for j in JUNIOR_SIGNALS)
    if has_senior and not has_junior:
        return "Senior+"
    if has_junior and not has_senior:
        return "Junior/Associate"
    return "Mid-level"

def extract_years_required(text: str) -> int:
    yrs = 0
    for m in RE_YEARS.finditer(text):
        try:
            yrs = max(yrs, int(m.group(1)))
        except Exception:
            pass
    return yrs

def split_required_vs_nice(jd_text: str) -> tuple[set, set]:
    lines = [l.strip() for l in jd_text.splitlines() if l.strip()]
    required_block, preferred_block, current = [], [], None
    for ln in lines:
        if RE_REQUIRED.search(ln):
            current = "req"; continue
        if RE_NICE.search(ln):
            current = "nice"; continue
        if current == "req":
            required_block.append(ln)
        elif current == "nice":
            preferred_block.append(ln)
    req_keys = set(extract_keywords("\n".join(required_block)))
    nice_keys = set(extract_keywords("\n".join(preferred_block)))
    if not req_keys:
        req_keys = set(extract_keywords(jd_text))
    return req_keys, nice_keys

def weighted_fit_score(resume_keys: set, req_keys: set, nice_keys: set) -> tuple[float, dict]:
    req_tech, req_soft = req_keys & TECH_KEYWORDS, req_keys & SOFT_SKILLS
    nice_tech, nice_soft = nice_keys & TECH_KEYWORDS, nice_keys & SOFT_SKILLS

    have_req_tech  = resume_keys & req_tech
    have_req_soft  = resume_keys & req_soft
    have_nice_tech = resume_keys & nice_tech
    have_nice_soft = resume_keys & nice_soft

    w_req_tech, w_req_soft, w_nice_tech, w_nice_soft = 45, 25, 20, 10
    sc_req_tech  = (len(have_req_tech)  / max(1, len(req_tech)))  * w_req_tech
    sc_req_soft  = (len(have_req_soft)  / max(1, len(req_soft)))  * w_req_soft
    sc_nice_tech = (len(have_nice_tech) / max(1, len(nice_tech))) * w_nice_tech
    sc_nice_soft = (len(have_nice_soft) / max(1, len(nice_soft))) * w_nice_soft

    score = round(min(100.0, sc_req_tech + sc_req_soft + sc_nice_tech + sc_nice_soft), 1)
    details = {
        "have_req_tech": sorted(have_req_tech),
        "miss_req_tech": sorted(req_tech - resume_keys),
        "have_req_soft": sorted(have_req_soft),
        "miss_req_soft": sorted(req_soft - resume_keys),
        "have_nice_tech": sorted(have_nice_tech),
        "miss_nice_tech": sorted(nice_tech - resume_keys),
        "have_nice_soft": sorted(have_nice_soft),
        "miss_nice_soft": sorted(nice_soft - resume_keys),
    }
    return score, details

# ---------------- UI: Login ----------------
def login_panel():
    st.sidebar.markdown("### Login")
    u = st.sidebar.text_input("Username", key="login_u")
    p = st.sidebar.text_input("Password", type="password", key="login_p")
    if st.sidebar.button("Login"):
        if authenticate(u, p):
            st.session_state.auth = {"logged_in": True, "user": u}
            st.success(f"Welcome {u}!")
        else:
            st.error("Invalid username or password.")

def logout_button():
    if st.sidebar.button("Logout"):
        st.session_state.auth = {"logged_in": False, "user": None}
        st.success("Logged out.")

# ---------------- Pages ----------------
def page_generate_summary():
    st.subheader("✍️ Generate Resume Summary")
    col1, col2 = st.columns(2)
    with col1:
        full_name = st.text_input("Your Full Name")
        role = st.text_input("Target Role / Job Title")
    with col2:
        experience = st.text_area("Work Experience (free text)", height=140)
        skills = st.text_area("Skills / Tools / Technologies", height=140)

    use_gpt = st.checkbox("Use GPT (if enabled)", value=False and USE_GPT)
    if st.button("Generate Summary"):
        if use_gpt and USE_GPT and client:
            prompt = f"Write a concise, ATS-friendly professional summary for {full_name} targeting {role}. Experience: {experience}. Skills: {skills}. 3-5 bullet highlights."
            out = gpt_chat(prompt)
        else:
            out = summary_offline(full_name, role, experience, skills)

        st.success("Summary generated!")
        st.markdown(out)

        DB["metrics"]["summaries"] += 1
        save_db(DB)
        st.session_state.metrics = DB["metrics"]

        st.download_button("Download .txt", out.encode("utf-8"), file_name="resume_summary.txt")
        st.download_button("Download .pdf", export_pdf(out), file_name="resume_summary.pdf", mime="application/pdf")
        st.download_button("Download .docx", export_docx(out), file_name="resume_summary.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # Template shortcut inside Summary page (optional)
    st.divider()
    st.caption("Or use a template:")
    tpl2 = st.selectbox("Template", ["(None)"] + list(PROMPT_TEMPLATES.keys()))
    tone2 = st.selectbox("Template Tone", PROMPT_TONES, index=1)
    if st.button("Generate with Template"):
        if tpl2 != "(None)":
            prompt = PROMPT_TEMPLATES[tpl2].format(
                role=role or "Target Role",
                skills=skills or "Python, SQL, AWS",
                experience=experience or "3-5 years in cross-functional teams delivering data products",
                tone=tone2
            )
            if USE_GPT and client:
                out = gpt_chat(prompt)
            else:
                out = f"(Offline mock)\n\nTemplate: {tpl2}\nTone: {tone2}\n\n" + summary_offline(
                    full_name=full_name or "Candidate", role=role, experience=experience, skills=skills
                )
            st.success("Template output:")
            st.markdown(out)
            st.download_button("Download .txt", out.encode("utf-8"), file_name="summary_template.txt")

def page_upload_resume():
    st.subheader("📤 Upload Resume & Generate Interview Questions")
    pdf = st.file_uploader("Upload Resume (PDF/DOCX/TXT)", type=["pdf","docx","txt"])
    if not pdf:
        return
    text = extract_text_generic(pdf)
    st.text_area("Extracted Text", text, height=220)
    qtype = st.selectbox("Question Type", ["Behavioral", "Technical", "Mixed"])
    count = st.slider("Number of Questions", 1, 10, 5)

    use_gpt = st.checkbox("Use GPT (if enabled)", value=False and USE_GPT)
    if st.button("Generate Questions"):
        if use_gpt and USE_GPT and client:
            prompt = f"Generate {count} {qtype} interview questions tailored to the following resume content:\n{text}"
            out = gpt_chat(prompt)
            qs = [x.strip("- ").strip() for x in out.split("\n") if x.strip()][:count]
        else:
            qs = questions_offline(text, qtype, count)

        st.success("Questions:")
        for i, q in enumerate(qs, 1):
            st.markdown(f"**{i}.** {q}")

        DB["metrics"]["resumes"] += 1
        DB["metrics"]["questions"] += len(qs)
        save_db(DB)
        st.session_state.metrics = DB["metrics"]

        dl = "\n".join([f"{i}. {q}" for i, q in enumerate(qs, 1)])
        st.download_button("Download Questions (.txt)", dl.encode("utf-8"), file_name="interview_questions.txt")

def page_prompt_lab():
    st.subheader("🧪 Prompt Lab — Templates & Customization")
    col1, col2 = st.columns(2)
    with col1:
        tpl = st.selectbox("Template", list(PROMPT_TEMPLATES.keys()))
        tone = st.selectbox("Tone", PROMPT_TONES, index=1)
        role = st.text_input("Target Role / Job Title")
    with col2:
        skills = st.text_area("Skills (comma or lines)", height=120)
        experience = st.text_area("Experience / Context (free text)", height=120)

    use_gpt = st.checkbox("Use GPT (if enabled)", value=False and USE_GPT)
    if st.button("Generate from Template"):
        prompt = PROMPT_TEMPLATES[tpl].format(
            role=role or "Target Role",
            skills=skills or "Python, SQL, AWS",
            experience=experience or "3-5 years in cross-functional teams delivering data products",
            tone=tone
        )
        if use_gpt and USE_GPT and client:
            out = gpt_chat(prompt)
        else:
            out = f"(Offline mock)\n\nTemplate: {tpl}\nTone: {tone}\n\n" + summary_offline(
                full_name="Candidate", role=role, experience=experience, skills=skills
            )

        st.success("Generated:")
        st.markdown(out)

        st.download_button("Download .txt", out.encode("utf-8"),
                           file_name="prompt_lab_output.txt", mime="text/plain")
        st.download_button("Download .pdf", export_pdf(out),
                           file_name="prompt_lab_output.pdf", mime="application/pdf")
        st.download_button("Download .docx", export_docx(out),
                           file_name="prompt_lab_output.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def page_job_fit_salary():
    st.subheader("🧭 Job Fit & Salary Alignment Analyzer (Weighted, Smarter)")
    c1, c2 = st.columns(2)
    with c1:
        jd_file = st.file_uploader("Job Description (PDF/DOCX/TXT)", type=["pdf","docx","txt"], key="jd_file")
        jd_text = st.text_area("…or paste JD text", height=200)
        if jd_file and not jd_text.strip():
            jd_text = extract_text_generic(jd_file)
    with c2:
        resume_file = st.file_uploader("Your Resume (PDF/DOCX/TXT)", type=["pdf","docx","txt"], key="resume_file")
        resume_text = st.text_area("…or paste resume text", height=200)
        if resume_file and not resume_text.strip():
            resume_text = extract_text_generic(resume_file)

    st.divider()
    role = st.text_input("Target Role (e.g., Data Scientist)")
    location_level = st.selectbox("Location Cost Tier", ["standard","high-cost","low-cost","remote"])
    expected_salary = st.number_input("Your Expected Salary (USD, annual)", min_value=30000, max_value=500000, step=1000)

    if st.button("Analyze"):
        if not jd_text.strip() or not resume_text.strip():
            st.error("Please provide both a JD and resume (upload or paste).")
            return

        jd_seniority = infer_seniority(jd_text)
        years_req = extract_years_required(jd_text)
        req_keys, nice_keys = split_required_vs_nice(jd_text)
        rs_keys = set(extract_keywords(resume_text))

        fit_score, detail = weighted_fit_score(rs_keys, req_keys, nice_keys)

        st.success(f"Weighted Fit Score: **{fit_score}%**")
        st.caption(f"JD Seniority: **{jd_seniority}**  |  Years indicated: **{years_req or 'n/a'}**")

        colA, colB = st.columns(2)
        with colA:
            st.markdown("**Required — You Have**")
            st.write(", ".join(detail["have_req_tech"] + detail["have_req_soft"]) or "—")
            st.markdown("**Nice-to-have — You Have**")
            st.write(", ".join(detail["have_nice_tech"] + detail["have_nice_soft"]) or "—")
        with colB:
            st.markdown("**Required — Missing (address these first)**")
            st.write(", ".join(detail["miss_req_tech"] + detail["miss_req_soft"]) or "—")
            st.markdown("**Nice-to-have — Missing**")
            st.write(", ".join(detail["miss_nice_tech"] + detail["miss_nice_soft"]) or "—")

        band = estimate_salary_band(role, location_level)
        if band:
            s = compare_salary(int(expected_salary), band)
            st.info(f"**Salary Alignment:** {s['status']}")
            st.caption(f"Market band: ${s['band_low']:,}–${s['band_high']:,} (mid ${s['band_mid']:,})")
            st.write(s["note"])
        else:
            st.warning("No salary band found for this role yet. Try a common title (e.g., 'Data Scientist').")

        recs = []
        if detail["miss_req_tech"] or detail["miss_req_soft"]:
            recs.append("Prioritize adding evidence for **required** missing items first.")
        if any(m in {"aws","azure","gcp"} for m in detail["miss_req_tech"] + detail["miss_nice_tech"]):
            recs.append("Consider an Associate-level cloud cert to quickly close credibility gaps.")
        if jd_seniority == "Senior+" and years_req and years_req > 6:
            recs.append("Emphasize scope/impact: systems scale, budgets, uptime, mentoring, cross-team influence.")
        if fit_score < 80:
            recs.append("Tailor summary to include 3–5 **required** JD keywords you already match.")
        if band and s["status"] == "Above Market":
            recs.append("Lower ask by ~10–15% or explicitly justify senior scope (team size, budgets, SLAs).")
        if not recs:
            recs.append("You're well aligned. Focus on STAR stories and results.")

        st.markdown("### Recommendations")
        for r in recs:
            st.markdown(f"- {r}")

        report = (
            f"ResumeReadyPro — Job Fit & Salary Alignment Report (Weighted)\n"
            f"Generated: {datetime.utcnow().isoformat()}Z\n\n"
            f"Role: {role or '(unspecified)'}  |  Location: {location_level}\n"
            f"Expected salary: ${int(expected_salary):,}\n"
            f"JD Seniority: {jd_seniority}  |  Years indicated: {years_req or 'n/a'}\n"
            f"Weighted Fit Score: {fit_score}%\n\n"
            f"Required (have): {', '.join(detail['have_req_tech'] + detail['have_req_soft']) or '—'}\n"
            f"Required (missing): {', '.join(detail['miss_req_tech'] + detail['miss_req_soft']) or '—'}\n"
            f"Nice-to-have (have): {', '.join(detail['have_nice_tech'] + detail['have_nice_soft']) or '—'}\n"
            f"Nice-to-have (missing): {', '.join(detail['miss_nice_tech'] + detail['miss_nice_soft']) or '—'}\n\n"
        )
        if band:
            report += (
                f"Market Band: ${s['band_low']:,}–${s['band_high']:,} (mid ${s['band_mid']:,})\n"
                f"Salary Alignment: {s['status']}\n{s['note']}\n\n"
            )
        report += "Recommendations:\n" + "\n".join([f"- {x}" for x in recs])

        st.download_button("⬇️ Download Report (.txt)", report.encode("utf-8"),
                           file_name="job_fit_salary_report.txt", mime="text/plain")
        st.download_button("⬇️ Download Report (.pdf)", export_pdf(report),
                           file_name="job_fit_salary_report.pdf", mime="application/pdf")
        st.download_button("⬇️ Download Report (.docx)", export_docx(report),
                           file_name="job_fit_salary_report.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        DB["metrics"]["gap_analyses"] += 1
        save_db(DB)
        st.session_state.metrics = DB["metrics"]

def page_admin():
    st.subheader("📊 Admin Dashboard")
    metrics = DB.get("metrics", {})
    st.write("**Totals**")
    st.json(metrics)

    users = DB.get("users", {})
    if users and pd is not None:
        df = pd.DataFrame([{"username": u, "name": info.get("name"), "created": info.get("created")} for u, info in users.items()])
        st.dataframe(df, use_container_width=True)
        csv = df.to_csv(index=False).encode("utf-8")
    else:
        df, csv = None, b""

    c1, c2 = st.columns(2)
    with c1:
        st.download_button("Download metrics.json", json.dumps(metrics, indent=2).encode("utf-8"), file_name="metrics.json")
    with c2:
        st.download_button("Download users.csv", csv, file_name="users.csv")

    if pd is not None and plt is not None and users:
        st.markdown("### Usage Chart")
        totals = pd.Series(metrics)
        fig, ax = plt.subplots()
        totals.plot(kind="bar", ax=ax)
        ax.set_title("ResumeReadyPro Usage Metrics")
        st.pyplot(fig)

def page_register():
    st.subheader("👤 Register New User")
    u = st.text_input("Username (lowercase)")
    n = st.text_input("Full Name")
    p = st.text_input("Password", type="password")
    if st.button("Register"):
        ok, msg = register_user(u, n, p)
        st.success(msg) if ok else st.error(msg)

def page_change_password():
    st.subheader("🔑 Change Password")
    if not st.session_state.auth.get("logged_in"):
        st.info("Log in to change your password.")
        return
    u = st.session_state.auth["user"]
    old = st.text_input("Old Password", type="password")
    new = st.text_input("New Password", type="password")
    if st.button("Update Password"):
        ok, msg = change_password(u, old, new)
        st.success(msg) if ok else st.error(msg)

def page_password_reset():
    st.subheader("🔒 Password Reset (Token-based)")
    tab1, tab2 = st.tabs(["Request Reset Token", "Reset with Token"])
    with tab1:
        uname = st.text_input("Username for reset")
        if st.button("Create Token"):
            ok, token_or_msg = create_reset_token(uname)
            if ok:
                st.success("Reset token created. Copy it now (no email in offline mode):")
                st.code(token_or_msg)
            else:
                st.error(token_or_msg)
    with tab2:
        uname = st.text_input("Username", key="rt_un")
        token = st.text_input("Reset Token", key="rt_tok")
        newp = st.text_input("New Password", type="password", key="rt_pw")
        if st.button("Reset Password"):
            ok, msg = reset_password_with_token(uname, token, newp)
            st.success(msg) if ok else st.error(msg)

def page_about():
    st.subheader("ℹ️ About ResumeReadyPro")
    st.markdown("""
**ResumeReadyPro** helps you:
- Generate resume summaries (offline or GPT when enabled)
- Extract text from resumes and practice interview questions
- Analyze Job Fit & Salary alignment with a clear action plan
- Explore templates & tone in the Prompt Lab
- Track usage via an Admin Dashboard and export your data

This hybrid build runs **offline** by default. Flip to GPT by setting `USE_GPT=True` and adding `OPENAI_API_KEY`.
""")

# ---------------- Navigation ----------------
def nav():
    return st.sidebar.radio("Go to", [
        "Generate Summary",
        "Upload Resume",
        "Job Fit & Salary Alignment",
        "Prompt Lab",
        "Admin Dashboard",
        "Register User",
        "Change Password",
        "Password Reset",
        "About"
    ])

# ---------------- Main ----------------
def main():
    st.set_page_config(page_title=APP_TITLE, page_icon="📄", layout="wide")
    st.title(APP_TITLE)
    st.caption("Secure offline build — ready to switch to GPT when you are.")

    # Onboarding
    onboarding_wizard()

    # Auth
    if not st.session_state.auth.get("logged_in"):
        st.sidebar.markdown("---")
        login_panel()
        st.stop()

    st.sidebar.markdown(f"**User:** {st.session_state.auth.get('user')}")
    logout_button()
    st.sidebar.markdown("---")

    page = nav()

    if page == "Generate Summary":
        page_generate_summary()
    elif page == "Upload Resume":
        page_upload_resume()
    elif page == "Job Fit & Salary Alignment":
        page_job_fit_salary()
    elif page == "Prompt Lab":
        page_prompt_lab()
    elif page == "Admin Dashboard":
        page_admin()
    elif page == "Register User":
        page_register()
    elif page == "Change Password":
        page_change_password()
    elif page == "Password Reset":
        page_password_reset()
    elif page == "About":
        page_about()
    else:
        st.info("Select a page from the sidebar.")

if __name__ == "__main__":
    main()
